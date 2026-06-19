#!/usr/bin/env python3
"""Generate the three DVAS Markdown deliverables as DOCX files.

The converter keeps wide Markdown tables readable by representing each row as a
label/detail field block. That transformation is deliberately silent in the
DOCX body: the generated documents should contain only source document content.
"""

from __future__ import annotations

import argparse
import hashlib
import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BODY_FONT = "Calibri"
CJK_FONT = "Microsoft YaHei"
ACCENT = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
TABLE_WIDTH_DXA = 9360


@dataclass(frozen=True)
class Deliverable:
    source: str
    output: str


DELIVERABLES = (
    Deliverable(
        source="数据收益分配系统_V1.3_需求规格说明书_导航结构更新版.md",
        output="数据收益分配系统_V1.3_需求规格说明书_导航结构更新版.docx",
    ),
    Deliverable(
        source="数据收益分配系统_系统详细功能设计_V1.1_导航结构更新版.md",
        output="数据收益分配系统_系统详细功能设计_V1.1_导航结构更新版.docx",
    ),
    Deliverable(
        source="数据收益分配系统_数据库设计与ER关系图_V1.0_导航结构更新版.md",
        output="数据收益分配系统_数据库设计与ER关系图_V1.0_导航结构更新版.docx",
    ),
)


def set_run_font(run, size=None, bold=None, italic=None, color=None, name=BODY_FONT):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), CJK_FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def clean_inline(text: str) -> str:
    text = text.replace("\\>", ">").replace("\\|", "|")
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.replace("**", "").replace("__", "").strip()


def add_runs_with_basic_bold(paragraph, text: str, size=11, color=None, italic=False):
    for part in re.split(r"(\*\*[^*]+\*\*)", text):
        if not part:
            continue
        bold = part.startswith("**") and part.endswith("**")
        value = part[2:-2] if bold else part
        value = re.sub(r"`([^`]+)`", r"\1", value)
        run = paragraph.add_run(value)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [clean_inline(cell) for cell in re.split(r"(?<!\\)\|", stripped)]


def is_separator_row(line: str) -> bool:
    cells = parse_table_row(line)
    return bool(cells) and all(
        re.fullmatch(r"\s*:?-{3,}:?\s*", cell or "") for cell in cells
    )


def apply_doc_styles(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_styles = (
        ("Heading 1", 16, ACCENT, 18, 10),
        ("Heading 2", 13, ACCENT, 14, 7),
        ("Heading 3", 12, DARK, 10, 5),
        ("Heading 4", 11, DARK, 8, 4),
    )
    for name, size, color, before, after in heading_styles:
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Code Block" not in [style.name for style in styles]:
        code = styles.add_style("Code Block", 1)
        code.font.name = "Courier New"
        code._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
        code.font.size = Pt(8.5)
        code.paragraph_format.space_after = Pt(3)
        code.paragraph_format.line_spacing = 1.0


def add_title(doc: Document, title: str, source_name: str):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run_font(run, size=20, bold=True, color=ACCENT)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(10)
    run = meta.add_run(f"源文件: {source_name}    生成日期: {date.today().isoformat()}")
    set_run_font(run, size=9, color=MUTED)


def add_table(doc: Document, rows: list[list[str]]):
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[1:]
    cols = len(header)
    normalized = [
        row + [""] * (cols - len(row)) if len(row) < cols else row[:cols]
        for row in [header] + body
    ]

    max_lens = [max(len(row[index]) for row in normalized) for index in range(cols)]
    weights = [max(1.0, min(4.0, length / 12.0)) for length in max_lens]
    total = sum(weights)
    widths = [max(720, int(TABLE_WIDTH_DXA * weight / total)) for weight in weights]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)

    table = doc.add_table(rows=1, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False

    for index, text in enumerate(header):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, widths[index])
        set_cell_shading(cell, "E8EEF5")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        add_runs_with_basic_bold(paragraph, text, size=9)
        for run in paragraph.runs:
            run.bold = True

    for row in body:
        cells = table.add_row().cells
        for index, text in enumerate(row):
            cell = cells[index]
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            add_runs_with_basic_bold(paragraph, text, size=8.5)
    doc.add_paragraph()


def add_wide_table_blocks(doc: Document, rows: list[list[str]]):
    if len(rows) < 2:
        return
    header = rows[0]
    body = rows[1:]
    for row_number, row in enumerate(body, start=1):
        normalized = (
            row + [""] * (len(header) - len(row))
            if len(row) < len(header)
            else row[: len(header)]
        )
        label = normalized[0] or f"第 {row_number} 行"
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.keep_with_next = True
        run = paragraph.add_run(f"{row_number}. {label}")
        set_run_font(run, size=10, bold=True, color=DARK)

        table = doc.add_table(rows=0, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        table.autofit = False
        for key, value in zip(header, normalized):
            if not (key or value):
                continue
            cells = table.add_row().cells
            for cell, width in zip(cells, (2100, 7260)):
                set_cell_width(cell, width)
                set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_shading(cells[0], "F2F4F7")
            left = cells[0].paragraphs[0]
            left.paragraph_format.space_after = Pt(0)
            add_runs_with_basic_bold(left, key, size=8.5)
            for run in left.runs:
                run.bold = True
            right = cells[1].paragraphs[0]
            right.paragraph_format.space_after = Pt(0)
            add_runs_with_basic_bold(right, value, size=8.5)

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)


def add_markdown_table(doc: Document, block: list[str]):
    rows = [parse_table_row(line) for line in block]
    if len(rows) >= 2 and is_separator_row(block[1]):
        rows = [rows[0]] + rows[2:]
    if not rows:
        return

    cols = max(len(row) for row in rows)
    max_cell = max((len(cell) for row in rows for cell in row), default=0)
    if cols <= 4 and max_cell <= 120 and len(rows) <= 80:
        add_table(doc, rows)
    else:
        add_wide_table_blocks(doc, rows)


def add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        paragraph = doc.add_paragraph(style="Code Block")
        run = paragraph.add_run(line.rstrip())
        set_run_font(run, size=8.5, name="Courier New")
        paragraph.paragraph_format.left_indent = Inches(0.12)


def source_title(lines: list[str], fallback: str) -> str:
    for line in lines:
        match = re.match(r"^#\s+(.+)$", line)
        if match:
            return clean_inline(match.group(1))
    return fallback


def convert_markdown(source: Path, output: Path):
    text = source.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = source_title(lines, source.stem)

    doc = Document()
    apply_doc_styles(doc)
    doc.core_properties.title = title
    doc.core_properties.subject = "DVAS documentation generated from Markdown source"
    add_title(doc, title, source.name)

    index = 0
    in_code = False
    code_lines: list[str] = []
    skipped_first_h1 = False

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue

        if not stripped:
            index += 1
            continue

        if stripped.startswith("|"):
            block = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                block.append(lines[index])
                index += 1
            add_markdown_table(doc, block)
            continue

        match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if match:
            level = min(len(match.group(1)), 4)
            heading_text = clean_inline(match.group(2))
            if level == 1 and heading_text == title and not skipped_first_h1:
                skipped_first_h1 = True
                index += 1
                continue
            doc.add_heading(heading_text, level=level)
            index += 1
            continue

        if stripped in {"---", "***", "___"}:
            index += 1
            continue

        if stripped.startswith(">"):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            add_runs_with_basic_bold(
                paragraph,
                clean_inline(stripped.lstrip(">").strip()),
                size=10,
                color=MUTED,
                italic=True,
            )
            index += 1
            continue

        bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            add_runs_with_basic_bold(paragraph, clean_inline(bullet.group(1)), size=11)
            index += 1
            continue

        numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            add_runs_with_basic_bold(paragraph, clean_inline(numbered.group(1)), size=11)
            index += 1
            continue

        paragraph = doc.add_paragraph()
        add_runs_with_basic_bold(paragraph, clean_inline(stripped), size=11)
        index += 1

    if in_code and code_lines:
        add_code_block(doc, code_lines)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


def generate(repo_root: Path, output_dir: Path) -> list[Path]:
    generated = []
    for deliverable in DELIVERABLES:
        source = repo_root / deliverable.source
        output = output_dir / deliverable.output
        convert_markdown(source, output)
        generated.append(output)
    return generated


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument(
        "--output-dir",
        default="docs/generated_docx",
        help="Destination folder for generated DOCX files.",
    )
    args = parser.parse_args()

    repo_root = Path.cwd()
    output_dir = repo_root / args.output_dir
    generated = generate(repo_root, output_dir)
    for path in generated:
        digest = hashlib.sha256(path.read_bytes()).hexdigest()[:12]
        print(f"generated {path.relative_to(repo_root)} sha256={digest}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
